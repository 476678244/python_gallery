"""Generate a single A4 page containing 6 score tables (班级/作业, 1-36) in a 2x3 grid."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, RGBColor

HERE = Path(__file__).parent
OUT = HERE / "打分表_6.docx"

# Make each table square: width == height
# 3 rows must fit the A4 usable height (297 - top/bottom margins),
# so vertical space is the binding constraint, not width.
TABLE_W_MM = 85          # each table is square; 3 stacked fill one page
HEADER_H_MM = 8
ROW_H_MM = (TABLE_W_MM - HEADER_H_MM) / 6  # 6 number rows -> total height == width
GAP_MM = 3               # white space between adjacent tables (both directions)


def set_cell_borders(cell, sz=4):
    """Add single black borders to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)


def set_no_borders(cell):
    """Remove borders from a cell (used for the outer layout table)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def set_row_height(row, mm, rule="atLeast"):
    trPr = row._tr.get_or_add_trPr()
    h = trPr.makeelement(qn("w:trHeight"), {})
    h.set(qn("w:val"), str(int(Mm(mm).twips)))
    h.set(qn("w:hRule"), rule)
    trPr.append(h)


def style_text(cell, text, bold=False, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_score_table(container_cell):
    """Build one 班级/作业 score table (header + numbers 1-36) inside a cell."""
    tbl = container_cell.add_table(rows=7, cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.allow_autofit = False

    # Fixed table layout with equal column widths so width is predictable
    tblPr = tbl._tbl.tblPr
    layout = tblPr.makeelement(qn("w:tblLayout"), {})
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    col_w = Mm(TABLE_W_MM / 6)
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = col_w

    # Header row: 班级 (cols 0-2) | 作业 (cols 3-5)
    hdr = tbl.rows[0]
    set_row_height(hdr, HEADER_H_MM)
    c0 = hdr.cells[0].merge(hdr.cells[1]).merge(hdr.cells[2])
    c1 = hdr.cells[3].merge(hdr.cells[4]).merge(hdr.cells[5])
    style_text(c0, "班级", bold=True, size=11)
    style_text(c1, "作业", bold=True, size=11)

    # Number rows 1-36
    n = 1
    for r in range(1, 7):
        row = tbl.rows[r]
        set_row_height(row, ROW_H_MM)
        for c in range(6):
            style_text(row.cells[c], str(n), size=15)
            n += 1

    # Borders on every cell
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_borders(cell)


def main():
    doc = Document()

    # A4 portrait, narrow margins to maximize space
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(5)
    sec.bottom_margin = Mm(5)
    sec.left_margin = Mm(8)
    sec.right_margin = Mm(8)

    # Outer 3x2 borderless layout table, sized to hug the tables (no extra
    # whitespace). Column width = table + gap so the only gap is GAP_MM.
    outer = doc.add_table(rows=3, cols=2)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    outer.allow_autofit = False

    tblPr = outer._tbl.tblPr
    layout = tblPr.makeelement(qn("w:tblLayout"), {})
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # zero cell margins so the column width is purely table + gap
    mar = tblPr.makeelement(qn("w:tblCellMar"), {})
    for edge in ("left", "right", "top", "bottom"):
        el = mar.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)

    col_w = Mm(TABLE_W_MM + GAP_MM)
    for r in range(3):
        # force each outer row to exactly table height + gap so rows hug the
        # tables instead of auto-expanding to fill the page
        set_row_height(outer.rows[r], TABLE_W_MM + GAP_MM, rule="exact")
        for c in range(2):
            cell = outer.cell(r, c)
            cell.width = col_w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_no_borders(cell)
            # drop the cell's default leading empty paragraph
            lead = cell.paragraphs[0]
            lead._element.getparent().remove(lead._element)
            build_score_table(cell)
            # Word requires a trailing paragraph after a table; keep it minimal
            tail = cell.add_paragraph()
            tail.paragraph_format.space_before = Pt(0)
            tail.paragraph_format.space_after = Pt(0)
            tail.paragraph_format.line_spacing = Pt(1)
            tail.add_run().font.size = Pt(1)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
